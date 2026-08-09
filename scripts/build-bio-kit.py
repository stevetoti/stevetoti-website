#!/usr/bin/env python3
"""Build branded PDF + DOCX versions of BIO-KIT.md."""
import re
import subprocess
import sys
from pathlib import Path

import markdown
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SRC = Path("/Users/stephentotimeh/Projects/stevetoti-website/BIO-KIT.md")
OUT = SRC.parent
NAVY = "1E3A8A"
GOLD = "F59E0B"
GREY = "6B7280"

md_text = SRC.read_text()

# ----------------------------------------------------------------- PDF (HTML → Chrome)
html_body = markdown.markdown(md_text, extensions=["tables", "fenced_code", "sane_lists"])
html = f"""<!doctype html><html><head><meta charset="utf-8"><title>Stephen Totimeh — Bio Kit</title>
<style>
@page {{ size: A4; margin: 18mm 16mm; }}
* {{ box-sizing: border-box; }}
body {{ font-family: -apple-system, "Helvetica Neue", Arial, sans-serif; color:#1f2937;
        font-size: 10.5pt; line-height: 1.55; margin:0; }}
h1 {{ color:#{NAVY}; font-size: 24pt; margin:0 0 4pt; letter-spacing:-.5pt; }}
h1 + p {{ color:#{GREY}; margin-top:0; }}
h2 {{ color:#{NAVY}; font-size: 15pt; margin:22pt 0 6pt; padding-bottom:4pt;
      border-bottom:2px solid #{GOLD}; page-break-after:avoid; }}
h3 {{ color:#{NAVY}; font-size: 11.5pt; margin:14pt 0 4pt; page-break-after:avoid; }}
p, li {{ orphans:2; widows:2; }}
code, pre {{ font-family: "SF Mono", Menlo, Consolas, monospace; }}
pre {{ background:#F8F7F4; border:1px solid #E5E7EB; border-left:3px solid #{GOLD};
       border-radius:4px; padding:9pt 11pt; font-size:9pt; line-height:1.5;
       white-space:pre-wrap; word-wrap:break-word; page-break-inside:avoid; }}
blockquote {{ margin:8pt 0; padding:6pt 12pt; background:#FEF6E7; border-left:3px solid #{GOLD};
              color:#4b5563; font-size:9.5pt; }}
blockquote p {{ margin:0; }}
table {{ border-collapse:collapse; width:100%; font-size:9.5pt; margin:8pt 0;
         page-break-inside:avoid; }}
th {{ background:#{NAVY}; color:#fff; text-align:left; padding:5pt 8pt; font-weight:600; }}
td {{ border-bottom:1px solid #E5E7EB; padding:5pt 8pt; vertical-align:top; }}
td:first-child {{ width:26%; font-weight:600; color:#{NAVY}; }}
hr {{ border:none; border-top:1px solid #E5E7EB; margin:16pt 0; }}
strong {{ color:#111827; }}
a {{ color:#{NAVY}; }}
</style></head><body>{html_body}
<p style="margin-top:22pt;color:#{GREY};font-size:8.5pt;border-top:1px solid #E5E7EB;padding-top:6pt;">
Stephen Totimeh · stevetoti.com · steve@pacificwavedigital.com · Generated from BIO-KIT.md</p>
</body></html>"""

tmp_html = Path("/tmp/bio_kit_render.html")
tmp_html.write_text(html)

chrome = "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome"
pdf_path = OUT / "Stephen-Totimeh-Bio-Kit.pdf"
ok = False
for flag in ("--headless=new", "--headless"):
    r = subprocess.run(
        [chrome, flag, "--disable-gpu", "--no-sandbox", "--no-pdf-header-footer",
         f"--print-to-pdf={pdf_path}", f"file://{tmp_html}"],
        capture_output=True, timeout=120,
    )
    if pdf_path.exists() and pdf_path.stat().st_size > 5000:
        ok = True
        break
if not ok:
    print("PDF FAILED", r.stderr.decode()[:400], file=sys.stderr)

# ----------------------------------------------------------------- DOCX
doc = Document()
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = Inches(0.8)
sec.left_margin = sec.right_margin = Inches(0.75)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)


def shade(par, hexcolor):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    par._p.get_or_add_pPr().append(el)


def rich(par, text, size=10.5, color=None, bold=False, mono=False):
    """Add text to a paragraph, honouring **bold** and `code` markers."""
    for chunk in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
        if not chunk:
            continue
        run = par.add_run(chunk.strip("*`"))
        run.bold = bold or chunk.startswith("**")
        run.font.size = Pt(size)
        run.font.name = "Consolas" if mono else "Calibri"
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


lines = md_text.split("\n")
i = 0
in_code = False
code_buf = []
while i < len(lines):
    line = lines[i]

    # fenced code blocks -> shaded copy-paste box
    if line.startswith("```"):
        if in_code:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.12)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(10)
            shade(p, "F8F7F4")
            for n, cl in enumerate(code_buf):
                if n:
                    p.add_run("\n")
                r = p.add_run(cl)
                r.font.name = "Consolas"
                r.font.size = Pt(9)
            code_buf, in_code = [], False
        else:
            in_code = True
        i += 1
        continue
    if in_code:
        code_buf.append(line)
        i += 1
        continue

    # tables
    if line.startswith("|") and i + 1 < len(lines) and set(lines[i + 1].replace("|", "").strip()) <= {"-", ":", " "}:
        rows = []
        j = i
        while j < len(lines) and lines[j].startswith("|"):
            cells = [c.strip() for c in lines[j].strip().strip("|").split("|")]
            if not set("".join(cells)) <= {"-", ":", " "}:
                rows.append(cells)
            j += 1
        if rows:
            t = doc.add_table(rows=0, cols=max(len(r) for r in rows))
            t.style = "Light Grid Accent 1"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ri, row in enumerate(rows):
                cells = t.add_row().cells
                for ci, val in enumerate(row):
                    if ci < len(cells):
                        cp = cells[ci].paragraphs[0]
                        rich(cp, val, size=9.5, bold=(ri == 0 or ci == 0),
                             color=NAVY if ci == 0 else None)
            doc.add_paragraph()
        i = j
        continue

    stripped = line.strip()
    if not stripped:
        i += 1
        continue

    if stripped.startswith("# "):
        h = doc.add_heading(level=0)
        r = h.add_run(stripped[2:])
        r.font.color.rgb = RGBColor.from_string(NAVY)
        r.font.size = Pt(24)
        r.font.name = "Calibri"
    elif stripped.startswith("## "):
        h = doc.add_heading(level=1)
        r = h.add_run(stripped[3:])
        r.font.color.rgb = RGBColor.from_string(NAVY)
        r.font.size = Pt(15)
        r.font.name = "Calibri"
    elif stripped.startswith("### "):
        h = doc.add_heading(level=2)
        r = h.add_run(stripped[4:])
        r.font.color.rgb = RGBColor.from_string(NAVY)
        r.font.size = Pt(11.5)
        r.font.name = "Calibri"
    elif stripped.startswith("> "):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        shade(p, "FEF6E7")
        rich(p, stripped[2:], size=9.5, color="4B5563")
    elif stripped.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        rich(p, stripped[2:])
    elif set(stripped) <= {"-"} and len(stripped) >= 3:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rich(p, stripped)
    i += 1

foot = doc.add_paragraph()
rich(foot, "Stephen Totimeh · stevetoti.com · steve@pacificwavedigital.com", size=8.5, color=GREY)

docx_path = OUT / "Stephen-Totimeh-Bio-Kit.docx"
doc.save(docx_path)

for f in (pdf_path, docx_path):
    print(f"{f.name}: {'OK' if f.exists() else 'MISSING'} ({f.stat().st_size // 1024} KB)" if f.exists() else f"{f.name}: MISSING")
